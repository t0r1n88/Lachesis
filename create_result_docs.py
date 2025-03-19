# -*- coding: UTF-8 -*-
"""
Скрипт для генерации документов по результатам тестирования
"""
import pandas as pd
pd.options.mode.copy_on_write = True
import os
import re
import time
import tempfile
from tkinter import messagebox
from docxcompose.composer import Composer
from docx import Document
from docxtpl import DocxTemplate
from docx2pdf import convert
import zipfile
import platform

class NotNumberColumn(Exception):
    """
    Исключение для обработки варианта когда в таблице нет колонки с таким порядковым номером
    """
    pass

class NoMoreNumberColumn(Exception):
    """
    Исключение для обработки варианта если введено больше 3 порядковых номеров колонок для создания структуры папок или
    больше 2 порядковых номеров для создания структуры названия файла
    """
    pass

class PdfLinux(Exception):
    """
    Исключение для контроля возможности запуска скрипта с опцией создания PDF в Linux
    """
    pass

class PDFAndFull(Exception):
    """
    Исключение для проверки нажатия обоих чекбоксов pdf и создания полного комплекта
    """
    pass

def combine_all_docx(data:str,name_file_template_doc,finish_path:str,mode_pdf:str, name_os:str):
    """
    Функция для объединения файлов Word взято отсюда
    https://stackoverflow.com/questions/24872527/combine-word-document-using-python-docx
    :param data: таблица с данными
    :param name_file_template_doc: файл шаблона
    :param finish_path: куда сохранять файл
    :param mode_pdf: создавать ли pdf
    :param name_os: для контроля операционной системы поскольку создание pdf работает только в Windows
    """
    # Список с созданными файлами
    files_lst = []
    # Создаем временную папку
    with tempfile.TemporaryDirectory() as tmpdirname:
        # Создаем и сохраняем во временную папку созданные документы Word
        for idx, row in enumerate(data):
            doc = DocxTemplate(name_file_template_doc)
            context = row
            doc.render(context)
            # Сохраняем файл
            doc.save(f'{tmpdirname}/{idx}.docx')
            # Добавляем путь\ к файлу в список
            files_lst.append(f'{tmpdirname}/{idx}.docx')
        # Получаем базовый файл
        main_doc = files_lst.pop(0)
        # Запускаем функцию

        number_of_sections = len(files_lst)
        # Открываем и обрабатываем базовый файл
        master = Document(main_doc)
        composer = Composer(master)
        # Перебираем и добавляем файлы к базовому
        for i in range(0, number_of_sections):
            doc_temp = Document(files_lst[i])
            composer.append(doc_temp)
        # Сохраняем файл
        composer.save(f"{finish_path}/Объединенный файл.docx")
        if mode_pdf == 'Yes':
            if name_os == 'Windows':
                if not os.path.exists(f'{finish_path}/PDF'):
                    os.makedirs(f'{finish_path}/PDF')
                convert(f'{finish_path}/Объединенный файл.docx', f'{finish_path}/PDF/Объединенный файл.pdf',
                        keep_active=True)




def prepare_entry_str(raw_str:str,pattern:str,repl_str:str,sep_lst:str)->list:
    """
    Функция для очистки строки от лишних символов и уменьшения на единицу (для нумерации с нуля)
    :param raw_str: обрабатываемая строка
    :param pattern: паттерн для замены символов
    :param repl_str: строка на которую нужно заменять символы
    :param sep_lst: разделитель по которому будет делиться список
    :return: список
    """
    raw_str = str(raw_str).replace('.',',')
    number_column_folder_structure = re.sub(pattern,repl_str,raw_str) # убираем из строки все лишние символы
    lst_number_column_folder_structure = number_column_folder_structure.split(sep_lst) # создаем список по запятой
    # отбрасываем возможные лишние элементы из за лишних запятых
    lst_number_column_folder_structure = [value for value in lst_number_column_folder_structure if value]
    # заменяем 0 на единицу
    lst_number_column_folder_structure = ['1' if x == '0' else x for x in lst_number_column_folder_structure]
    # Превращаем в числа и отнимаем 1 чтобы соответствовать индексам питона
    lst_number_column_folder_structure = list(map(lambda x:int(x)-1,lst_number_column_folder_structure))

    return lst_number_column_folder_structure


def zip_folder(folder_name, output_filename):
    # Создаем zip-файл
    with zipfile.ZipFile(f'{folder_name}/{output_filename}', 'w', compression=zipfile.ZIP_DEFLATED
                         ) as ziph:
        # Берем содержимое папки
        for root, dirs, files in os.walk(folder_name):
            # Проходимся по каждому файлу
            for file in files:
                if not file.startswith('~') and file.endswith('.docx'):
                    # Создаем полный путь к файлу
                    src_path = os.path.join(root, file)
                    # Добавляем файл в zip-архив
                    ziph.write(src_path, arcname=os.path.relpath(src_path, folder_name))



def save_result_file(finish_path:str,name_file:str,doc:DocxTemplate,idx:int,mode_pdf:str,name_os:str):
    """
    Функция для сохранения результатов
    :param finish_path: путь к папке сохранения
    :param name_file: название файла
    :param doc: объект DocxTemplate
    :param idx: счетчик
    :param mode_pdf: чекбокс сохранения PDF
    :param name_os: для контроля операционной системы поскольку создание pdf работает только в Windows
    :return:
    """
    if os.path.exists(f'{finish_path}/{name_file}.docx'):
        doc.save(f'{finish_path}/{name_file}_{idx}.docx')
        if mode_pdf == 'Yes':
            if name_os == 'Windows':
                if not os.path.exists(f'{finish_path}/PDF'):
                    os.makedirs(f'{finish_path}/PDF')
                convert(f'{finish_path}/{name_file}_{idx}.docx', f'{finish_path}/PDF/{name_file}_{idx}.pdf',
                        keep_active=True)
    else:
        doc.save(f'{finish_path}/{name_file}.docx')
        if mode_pdf == 'Yes':
            if name_os == 'Windows':
                if not os.path.exists(f'{finish_path}/PDF'):
                    os.makedirs(f'{finish_path}/PDF')
                convert(f'{finish_path}/{name_file}.docx', f'{finish_path}/PDF/{name_file}.pdf',
                        keep_active=True)

def short_version_save_result_file(finish_path:str,name_file:str,doc:DocxTemplate,idx:int):
    """
    Функция для сохранения результатов
    :param finish_path: путь к папке сохранения
    :param name_file: название файла
    :param doc: объект DocxTemplate
    :param idx: счетчик
    :return:
    """
    if os.path.exists(f'{finish_path}/{name_file}.docx'):
        doc.save(f'{finish_path}/{name_file}_{idx}.docx')
        convert(f'{finish_path}/{name_file}_{idx}.docx', f'{finish_path}/{name_file}_{idx}.pdf',
                keep_active=True)
        os.remove(f'{finish_path}/{name_file}_{idx}.docx')
    else:
        doc.save(f'{finish_path}/{name_file}.docx')
        convert(f'{finish_path}/{name_file}.docx', f'{finish_path}/{name_file}.pdf',
                keep_active=True)
        os.remove(f'{finish_path}/{name_file}.docx')







def generate_result_docs(name_file_data_doc:str,name_file_template_doc:str,path_to_end_folder_doc:str,
                         folder_structure:str,name_file:str,name_type_file:str,mode_pdf:str,mode_full:str):
    """
    Функция для генерации документов по результатам тестирования с разбиением по папкам и определенным названиям
    :param name_file_data_doc: таблица с яндекс форм
    :param name_file_template_doc: файл шаблона
    :param path_to_end_folder_doc: куда сохранять результаты
    :param folder_structure:строка с порядковыми номерами колонок по которым будет создаваться структура папок - 3,4
    :param name_file: строка с порядковыми номерами колонок по которым будет создаваться имя файла - 5,6,7
    :param name_type_file: название создаваемых документов которое будет находится в начале имен создаваемых файлов - Справка, Результат
    :param mode_pdf: нужно ли создавать пдф версии создаваемых файлов
    :param mode_full: по умолчанию создаются только pdf версии документов если включен то создаются docx и объединенные файлы
    :return:
    """
    try:
        name_os = platform.system()  # получаем платформу на которой запущена программа
        if name_os !='Windows' and mode_pdf =='Yes':
            raise PdfLinux

        if mode_full == 'Yes' and mode_pdf == 'Yes':
            raise PDFAndFull
        # Считываем данные
        # Добавил параметр dtype =str чтобы данные не преобразовались а использовались так как в таблице
        df = pd.read_excel(name_file_data_doc, dtype=str)
        df.fillna('Не заполнено',inplace=True)
        # очищаем строку от лишних символов и превращаем в список номеров колонок
        lst_number_column_folder_structure = prepare_entry_str(folder_structure,r'[^\d,]','',',')

        # Сохраняем без создания структуры папок
        if len(lst_number_column_folder_structure) == 0:
            # обрабатываем колонки с именем с названием файла
            # очищаем строку от лишних символов и превращаем в список номеров колонок
            lst_number_column_name_file = prepare_entry_str(name_file, r'[^\d,]', '', ',')

            # проверяем длину списка не более 2 и не равно 0
            if len(lst_number_column_name_file) == 0 or len(lst_number_column_name_file) > 2:
                raise NoMoreNumberColumn

            # проверяем чтобы номер колонки не превышал количество колонок в датафрейме
            for number_column in lst_number_column_name_file:
                if number_column > df.shape[1]:
                    raise NotNumberColumn

            # Добавляем разрыв в шаблон чтобы объединенный файл был без смешивания
            # Открываем шаблон
            doc_page_break = Document(name_file_template_doc)
            # Добавляем разрыв страницы
            doc_page_break.add_page_break()
            template_page_break_path = os.path.dirname(name_file_template_doc)
            # Сохраняем изменения в новом файле
            doc_page_break.save(f'{template_page_break_path}/page_break.docx')
            template_page_break_path_finish = f'{template_page_break_path}/page_break.docx'

            if mode_full == 'No':
                # переименовываем колонки указанные в качестве идентифицирующих для того чтобы они отображалисьвнутри файла
                if len(lst_number_column_name_file) == 1:
                    # если указана только одна колонка
                    name_column = df.columns[lst_number_column_name_file[0]]
                    df.rename(columns={name_column: 'Код_1'}, inplace=True)
                elif len(lst_number_column_name_file) == 2:
                    name_main_column = df.columns[lst_number_column_name_file[0]]  # первая колонка
                    name_second_column = df.columns[lst_number_column_name_file[1]]  # вторая колонка
                    df.rename(columns={name_main_column: 'Код_1', name_second_column: 'Код_2'}, inplace=True)

                data = df.to_dict('records')

                combine_all_docx(data, template_page_break_path_finish, path_to_end_folder_doc, mode_pdf, name_os)

                # В зависимости от состояния чекбоксов обрабатываем файлы
                # Создаем в цикле документы
                if len(lst_number_column_name_file) == 1:
                    # если указана только одна колонка
                    name_column = df.columns[lst_number_column_name_file[0]]
                    df.rename(columns={name_column: 'Код_1'}, inplace=True)

                    for idx, row in enumerate(data):
                        doc = DocxTemplate(name_file_template_doc)
                        context = row
                        doc.render(context)
                        name_file = f'{name_type_file} {row[name_column]}'
                        name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                        threshold_name = 200 - (len(path_to_end_folder_doc) + 10)
                        if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                            raise OSError
                        name_file = name_file[:threshold_name]  # ограничиваем название файла
                        # Сохраняем файл
                        save_result_file(path_to_end_folder_doc, name_file, doc, idx, mode_pdf, name_os)
                    zip_folder(path_to_end_folder_doc, f'Результаты тестирования.zip')


                elif len(lst_number_column_name_file) == 2:
                    name_main_column = df.columns[lst_number_column_name_file[0]]  # первая колонка
                    name_second_column = df.columns[lst_number_column_name_file[1]]  # вторая колонка
                    for idx, row in enumerate(data):
                        doc = DocxTemplate(name_file_template_doc)
                        context = row
                        doc.render(context)
                        # Сохраняем файл
                        name_file = f'{name_type_file} {row[name_main_column]} {row[name_second_column]}'
                        name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                        threshold_name = 200 - (len(path_to_end_folder_doc) + 10)
                        if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                            raise OSError
                        name_file = name_file[:threshold_name]  # ограничиваем название файла
                        # Сохраняем файл
                        save_result_file(path_to_end_folder_doc, name_file, doc, idx, mode_pdf, name_os)
                    zip_folder(path_to_end_folder_doc, f'Результаты тестирования.zip')
            else:
                # переименовываем колонки указанные в качестве идентифицирующих для того чтобы они отображалисьвнутри файла
                if len(lst_number_column_name_file) == 1:
                    # если указана только одна колонка
                    name_column = df.columns[lst_number_column_name_file[0]]
                    df.rename(columns={name_column: 'Код_1'}, inplace=True)
                elif len(lst_number_column_name_file) == 2:
                    name_main_column = df.columns[lst_number_column_name_file[0]]  # первая колонка
                    name_second_column = df.columns[lst_number_column_name_file[1]]  # вторая колонка
                    df.rename(columns={name_main_column: 'Код_1', name_second_column: 'Код_2'}, inplace=True)

                data = df.to_dict('records')
                # В зависимости от состояния чекбоксов обрабатываем файлы
                # Создаем в цикле документы
                if len(lst_number_column_name_file) == 1:
                    # если указана только одна колонка
                    name_column = df.columns[lst_number_column_name_file[0]]
                    df.rename(columns={name_column: 'Код_1'}, inplace=True)

                    for idx, row in enumerate(data):
                        doc = DocxTemplate(name_file_template_doc)
                        context = row
                        doc.render(context)
                        name_file = f'{name_type_file} {row[name_column]}'
                        name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                        threshold_name = 200 - (len(path_to_end_folder_doc) + 10)
                        if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                            raise OSError
                        name_file = name_file[:threshold_name]  # ограничиваем название файла
                        # Сохраняем файл
                        short_version_save_result_file(path_to_end_folder_doc, name_file, doc, idx)


                elif len(lst_number_column_name_file) == 2:
                    name_main_column = df.columns[lst_number_column_name_file[0]]  # первая колонка
                    name_second_column = df.columns[lst_number_column_name_file[1]]  # вторая колонка
                    for idx, row in enumerate(data):
                        doc = DocxTemplate(name_file_template_doc)
                        context = row
                        doc.render(context)
                        # Сохраняем файл
                        name_file = f'{name_type_file} {row[name_main_column]} {row[name_second_column]}'
                        name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                        threshold_name = 200 - (len(df) + 10)
                        if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                            raise OSError
                        name_file = name_file[:threshold_name]  # ограничиваем название файла
                        # Сохраняем файл
                        short_version_save_result_file(path_to_end_folder_doc, name_file, doc, idx)


        else:

            # проверяем длину списка не более 4
            if len(lst_number_column_folder_structure) > 4:
                raise NoMoreNumberColumn

            # проверяем чтобы номер колонки не превышал количество колонок в датафрейме
            for number_column in lst_number_column_folder_structure:
                if number_column > df.shape[1]:
                    raise NotNumberColumn

            # обрабатываем колонки с именем с названием файла
            # очищаем строку от лишних символов и превращаем в список номеров колонок
            lst_number_column_name_file = prepare_entry_str(name_file,r'[^\d,]','',',')

            # проверяем длину списка не более 2 и не равно 0
            if len(lst_number_column_name_file) == 0 or len(lst_number_column_name_file) > 2:
                raise NoMoreNumberColumn

            # проверяем чтобы номер колонки не превышал количество колонок в датафрейме
            for number_column in lst_number_column_name_file:
                if number_column > df.shape[1]:
                    raise NotNumberColumn

            # Добавляем разрыв в шаблон чтобы объединенный файл был без смешивания
            # Открываем шаблон
            doc_page_break = Document(name_file_template_doc)
            # Добавляем разрыв страницы
            doc_page_break.add_page_break()
            template_page_break_path = os.path.dirname(name_file_template_doc)
            # Сохраняем изменения в новом файле
            doc_page_break.save(f'{template_page_break_path}/page_break.docx')
            template_page_break_path_finish = f'{template_page_break_path}/page_break.docx'

            if mode_full == 'No':
                if len(lst_number_column_folder_structure) == 1:
                    # Если нужно создавать одноуровневую структуру
                    # получаем название колонки
                    main_layer_name_column = df.columns[lst_number_column_folder_structure[0]]
                    lst_unique_value = df[main_layer_name_column].unique() # получаем список уникальных значений
                    for name_folder in lst_unique_value:
                        temp_df = df[df[main_layer_name_column] == name_folder] # фильтруем по названию
                        # Конвертируем датафрейм в список словарей
                        clean_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_', name_folder)  # очищаем название от лишних символов
                        finish_path = f'{path_to_end_folder_doc}/{clean_name_folder}'
                        if not os.path.exists(finish_path):
                            os.makedirs(finish_path)

                        # переименовываем колонки указанные в качестве идентифицирующих для того чтобы они отображалисьвнутри файла
                        if len(lst_number_column_name_file) == 1:
                            # если указана только одна колонка
                            name_column = temp_df.columns[lst_number_column_name_file[0]]
                            temp_df.rename(columns={name_column: 'Код_1'}, inplace=True)
                        elif len(lst_number_column_name_file) == 2:
                            name_main_column = temp_df.columns[lst_number_column_name_file[0]] # первая колонка
                            name_second_column = temp_df.columns[lst_number_column_name_file[1]] # вторая колонка
                            temp_df.rename(columns={name_main_column: 'Код_1',name_second_column: 'Код_2'}, inplace=True)

                        data = temp_df.to_dict('records')


                        combine_all_docx(data,template_page_break_path_finish,finish_path,mode_pdf,name_os)

                        # В зависимости от состояния чекбоксов обрабатываем файлы
                        # Создаем в цикле документы
                        if len(lst_number_column_name_file) == 1:
                            # если указана только одна колонка
                            name_column = temp_df.columns[lst_number_column_name_file[0]]
                            temp_df.rename(columns={name_column: 'Код_1'}, inplace=True)


                            for idx, row in enumerate(data):
                                doc = DocxTemplate(name_file_template_doc)
                                context = row
                                doc.render(context)
                                name_file = f'{name_type_file} {row[name_column]}'
                                name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                threshold_name = 200 - (len(finish_path) + 10)
                                if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                    raise OSError
                                name_file = name_file[:threshold_name]  # ограничиваем название файла
                                # Сохраняем файл
                                save_result_file(finish_path,name_file,doc,idx,mode_pdf,name_os)
                            zip_folder(finish_path, f'Результаты тестирования {clean_name_folder}.zip')


                        elif len(lst_number_column_name_file) == 2:
                            name_main_column = temp_df.columns[lst_number_column_name_file[0]] # первая колонка
                            name_second_column = temp_df.columns[lst_number_column_name_file[1]] # вторая колонка
                            for idx, row in enumerate(data):
                                doc = DocxTemplate(name_file_template_doc)
                                context = row
                                doc.render(context)
                                # Сохраняем файл
                                name_file = f'{name_type_file} {row[name_main_column]} {row[name_second_column]}'
                                name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                threshold_name = 200 - (len(finish_path) + 10)
                                if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                    raise OSError
                                name_file = name_file[:threshold_name]  # ограничиваем название файла
                                # Сохраняем файл
                                save_result_file(finish_path,name_file,doc,idx,mode_pdf,name_os)
                            zip_folder(finish_path, f'Результаты тестирования {clean_name_folder}.zip')

                elif len(lst_number_column_folder_structure) == 2:
                    # Если нужно создавать двухуровневую структуру
                    # получаем название колонки для первого уровня папок
                    name_first_layer_column = df.columns[lst_number_column_folder_structure[0]]
                    name_second_layer_column = df.columns[lst_number_column_folder_structure[1]]

                    lst_unique_value_first_layer = df[name_first_layer_column].unique()  # получаем список уникальных значений
                    for first_name_folder in lst_unique_value_first_layer:
                        clean_first_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                         first_name_folder)  # очищаем название от лишних символов

                        # получаем отфильтрованный датафрейм по значениям колонки первого уровня
                        temp_df_first_layer = df[df[name_first_layer_column] == first_name_folder]  # фильтруем по названию
                        lst_unique_value_second_layer = temp_df_first_layer[name_second_layer_column].unique()  # получаем список уникальных значений
                        # фильтруем по значениям колонки второго уровня
                        for second_name_folder in lst_unique_value_second_layer:
                            temp_df_second_layer = temp_df_first_layer[temp_df_first_layer[name_second_layer_column] == second_name_folder]
                            clean_second_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_', second_name_folder)  # очищаем название от лишних символов

                            finish_path = f'{path_to_end_folder_doc}/{clean_first_name_folder}/{clean_second_name_folder}'
                            if not os.path.exists(finish_path):
                                os.makedirs(finish_path)

                            # переименовываем колонки указанные в качестве идентифицирующих для того чтобы они отображалисьвнутри файла
                            if len(lst_number_column_name_file) == 1:
                                # если указана только одна колонка
                                name_column = temp_df_second_layer.columns[lst_number_column_name_file[0]]
                                temp_df_second_layer.rename(columns={name_column: 'Код_1'}, inplace=True)
                            elif len(lst_number_column_name_file) == 2:
                                name_main_column = temp_df_second_layer.columns[lst_number_column_name_file[0]]  # первая колонка
                                name_second_column = temp_df_second_layer.columns[lst_number_column_name_file[1]]  # вторая колонка
                                temp_df_second_layer.rename(columns={name_main_column: 'Код_1', name_second_column: 'Код_2'}, inplace=True)

                            data = temp_df_second_layer.to_dict('records') # конвертируем в список словарей
                            # Создаем объединенный файл в формате docx и pdf
                            combine_all_docx(data, template_page_break_path_finish, finish_path, mode_pdf,name_os)
                            # Создаем в цикле документы
                            if len(lst_number_column_name_file) == 1:
                                # если указана только одна колонка
                                name_column = temp_df_second_layer.columns[lst_number_column_name_file[0]]
                                for idx, row in enumerate(data):
                                    doc = DocxTemplate(name_file_template_doc)
                                    context = row
                                    doc.render(context)
                                    # Сохраняем файл
                                    name_file = f'{name_type_file} {row[name_column]}'
                                    name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                    threshold_name = 200 - (len(finish_path) + 10)
                                    if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                        raise OSError
                                    name_file = name_file[:threshold_name]  # ограничиваем название файла
                                    # Сохраняем файл
                                    save_result_file(finish_path, name_file, doc, idx, mode_pdf,name_os)
                                zip_folder(finish_path, f'Результаты тестирования {clean_first_name_folder}_{clean_second_name_folder}.zip') # архивируем файлы docx


                            elif len(lst_number_column_name_file) == 2:
                                name_main_column = temp_df_second_layer.columns[lst_number_column_name_file[0]]  # первая колонка
                                name_second_column = temp_df_second_layer.columns[lst_number_column_name_file[1]]  # вторая колонка
                                for idx, row in enumerate(data):
                                    doc = DocxTemplate(name_file_template_doc)
                                    context = row
                                    doc.render(context)
                                    # Сохраняем файл
                                    name_file = f'{name_type_file} {row[name_main_column]} {row[name_second_column]}'
                                    name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                    threshold_name = 200 - (len(finish_path) + 10)
                                    if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                        raise OSError
                                    name_file = name_file[:threshold_name]  # ограничиваем название файла
                                    # Сохраняем файл
                                    save_result_file(finish_path, name_file, doc, idx, mode_pdf,name_os)
                                zip_folder(finish_path,f'Результаты тестирования {clean_first_name_folder}_{clean_second_name_folder}.zip')  # архивируем файлы docx


                elif len(lst_number_column_folder_structure) == 3:
                    # Если нужно создавать трехуровневую структуру Например Школа-Класс-буква класса
                    # получаем названия колонок для трех уровней
                    name_first_layer_column = df.columns[lst_number_column_folder_structure[0]]
                    name_second_layer_column = df.columns[lst_number_column_folder_structure[1]]
                    name_third_layer_column = df.columns[lst_number_column_folder_structure[2]]

                    lst_unique_value_first_layer = df[name_first_layer_column].unique()  # получаем список уникальных значений
                    for first_name_folder in lst_unique_value_first_layer:
                        clean_first_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                         first_name_folder)  # очищаем название от лишних символов

                        # получаем отфильтрованный датафрейм по значениям колонки первого уровня
                        temp_df_first_layer = df[df[name_first_layer_column] == first_name_folder]  # фильтруем по названию
                        lst_unique_value_second_layer = temp_df_first_layer[
                            name_second_layer_column].unique()  # получаем список уникальных значений второго уровня
                        # фильтруем по значениям колонки второго уровня
                        for second_name_folder in lst_unique_value_second_layer:
                            temp_df_second_layer = temp_df_first_layer[
                                temp_df_first_layer[name_second_layer_column] == second_name_folder]
                            clean_second_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                              second_name_folder)  # очищаем название от лишних символов
                            lst_unique_value_third_layer = temp_df_second_layer[
                                name_third_layer_column].unique()  # получаем список уникальных значений третьего уровня
                            for third_name_folder in lst_unique_value_third_layer:
                                clean_third_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                                  third_name_folder)  # очищаем название от лишних символов
                                temp_df_third_layer = temp_df_second_layer[
                                    temp_df_second_layer[name_third_layer_column] == third_name_folder]

                                finish_path = f'{path_to_end_folder_doc}/{clean_first_name_folder}/{clean_second_name_folder}/{clean_third_name_folder}'
                                if not os.path.exists(finish_path):
                                    os.makedirs(finish_path)
                                if len(lst_number_column_name_file) == 1:
                                    # если указана только одна колонка
                                    name_column = temp_df_third_layer.columns[lst_number_column_name_file[0]]
                                    temp_df_third_layer.rename(columns={name_column: 'Код_1'}, inplace=True)
                                elif len(lst_number_column_name_file) == 2:
                                    name_main_column = temp_df_third_layer.columns[lst_number_column_name_file[0]]  # первая колонка
                                    name_second_column = temp_df_third_layer.columns[lst_number_column_name_file[1]]  # вторая колонка
                                    temp_df_third_layer.rename(columns={name_main_column: 'Код_1', name_second_column: 'Код_2'}, inplace=True)

                                data = temp_df_third_layer.to_dict('records')  # конвертируем в список словарей
                                # Создаем объединенный файл в формате docx и pdf
                                combine_all_docx(data, template_page_break_path_finish, finish_path, mode_pdf,name_os)

                                # Создаем в цикле документы
                                if len(lst_number_column_name_file) == 1:
                                    # если указана только одна колонка
                                    name_column = temp_df_third_layer.columns[lst_number_column_name_file[0]]
                                    for idx, row in enumerate(data):
                                        doc = DocxTemplate(name_file_template_doc)
                                        context = row
                                        doc.render(context)
                                        # Сохраняем файл
                                        name_file = f'{name_type_file} {row[name_column]}'
                                        name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                        threshold_name = 200 - (len(finish_path) + 10)
                                        if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                            raise OSError
                                        name_file = name_file[:threshold_name]  # ограничиваем название файла
                                        # Сохраняем файл
                                        save_result_file(finish_path, name_file, doc, idx, mode_pdf,name_os)
                                    zip_folder(finish_path,f'Результаты тестирования {clean_first_name_folder}_{clean_second_name_folder}_{clean_third_name_folder}.zip')  # архивируем файлы docx



                                elif len(lst_number_column_name_file) == 2:
                                    name_main_column = temp_df_third_layer.columns[
                                        lst_number_column_name_file[0]]  # первая колонка
                                    name_second_column = temp_df_third_layer.columns[
                                        lst_number_column_name_file[1]]  # вторая колонка
                                    for idx, row in enumerate(data):
                                        doc = DocxTemplate(name_file_template_doc)
                                        context = row
                                        doc.render(context)
                                        # Сохраняем файл
                                        name_file = f'{name_type_file} {row[name_main_column]} {row[name_second_column]}'
                                        name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                        threshold_name = 200 - (len(finish_path) + 10)
                                        if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                            raise OSError
                                        name_file = name_file[:threshold_name]  # ограничиваем название файла
                                        # Сохраняем файл
                                        save_result_file(finish_path, name_file, doc, idx, mode_pdf,name_os)
                                    zip_folder(finish_path,
                                               f'Результаты тестирования {clean_first_name_folder}_{clean_second_name_folder}_{clean_third_name_folder}.zip')  # архивируем файлы docx
                elif len(lst_number_column_folder_structure) == 4:
                    # Если нужно создавать четырех уровневую структуру Например Результат- Число результата--Класс-буква класса
                    # получаем названия колонок для 4 уровней
                    name_first_layer_column = df.columns[lst_number_column_folder_structure[0]]
                    name_second_layer_column = df.columns[lst_number_column_folder_structure[1]]
                    name_third_layer_column = df.columns[lst_number_column_folder_structure[2]]
                    name_four_layer_column = df.columns[lst_number_column_folder_structure[3]]

                    lst_unique_value_first_layer = df[name_first_layer_column].unique()  # получаем список уникальных значений
                    for first_name_folder in lst_unique_value_first_layer:
                        clean_first_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                         first_name_folder)  # очищаем название от лишних символов

                        # получаем отфильтрованный датафрейм по значениям колонки первого уровня
                        temp_df_first_layer = df[df[name_first_layer_column] == first_name_folder]  # фильтруем по названию
                        lst_unique_value_second_layer = temp_df_first_layer[
                            name_second_layer_column].unique()  # получаем список уникальных значений второго уровня
                        # фильтруем по значениям колонки второго уровня
                        for second_name_folder in lst_unique_value_second_layer:
                            temp_df_second_layer = temp_df_first_layer[
                                temp_df_first_layer[name_second_layer_column] == second_name_folder]
                            clean_second_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                              second_name_folder)  # очищаем название от лишних символов
                            lst_unique_value_third_layer = temp_df_second_layer[
                                name_third_layer_column].unique()  # получаем список уникальных значений третьего уровня

                            for third_name_folder in lst_unique_value_third_layer:
                                clean_third_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                                  third_name_folder)  # очищаем название от лишних символов
                                temp_df_third_layer = temp_df_second_layer[
                                    temp_df_second_layer[name_third_layer_column] == third_name_folder]
                                lst_unique_value_four_layer = temp_df_third_layer[
                                    name_four_layer_column].unique()  # получаем список уникальных значений четвертого уровня

                                for four_name_folder in lst_unique_value_four_layer:
                                    clean_four_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                                     four_name_folder)  # очищаем название от лишних символов
                                    temp_df_four_layer = temp_df_third_layer[
                                        temp_df_third_layer[name_four_layer_column] == four_name_folder]


                                    finish_path = f'{path_to_end_folder_doc}/{clean_first_name_folder}/{clean_second_name_folder}/{clean_third_name_folder}/{clean_four_name_folder}'
                                    if not os.path.exists(finish_path):
                                        os.makedirs(finish_path)
                                    if len(lst_number_column_name_file) == 1:
                                        # если указана только одна колонка
                                        name_column = temp_df_third_layer.columns[lst_number_column_name_file[0]]
                                        temp_df_four_layer.rename(columns={name_column: 'Код_1'}, inplace=True)
                                    elif len(lst_number_column_name_file) == 2:
                                        name_main_column = temp_df_four_layer.columns[lst_number_column_name_file[0]]  # первая колонка
                                        name_second_column = temp_df_four_layer.columns[lst_number_column_name_file[1]]  # вторая колонка
                                        temp_df_four_layer.rename(columns={name_main_column: 'Код_1', name_second_column: 'Код_2'}, inplace=True)

                                    data = temp_df_four_layer.to_dict('records')  # конвертируем в список словарей
                                    # Создаем объединенный файл в формате docx и pdf
                                    combine_all_docx(data, template_page_break_path_finish, finish_path, mode_pdf,name_os)

                                    # Создаем в цикле документы
                                    if len(lst_number_column_name_file) == 1:
                                        # если указана только одна колонка
                                        name_column = temp_df_four_layer.columns[lst_number_column_name_file[0]]
                                        for idx, row in enumerate(data):
                                            doc = DocxTemplate(name_file_template_doc)
                                            context = row
                                            doc.render(context)
                                            # Сохраняем файл
                                            name_file = f'{name_type_file} {row[name_column]}'
                                            name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                            threshold_name = 200 - (len(finish_path) + 10)
                                            if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                                raise OSError
                                            name_file = name_file[:threshold_name]  # ограничиваем название файла
                                            # Сохраняем файл
                                            save_result_file(finish_path, name_file, doc, idx, mode_pdf,name_os)
                                        zip_folder(finish_path,f'Результаты тестирования {clean_first_name_folder}_{clean_second_name_folder}_{clean_third_name_folder}_{clean_four_name_folder}.zip')  # архивируем файлы docx



                                    elif len(lst_number_column_name_file) == 2:
                                        name_main_column = temp_df_four_layer.columns[
                                            lst_number_column_name_file[0]]  # первая колонка
                                        name_second_column = temp_df_four_layer.columns[
                                            lst_number_column_name_file[1]]  # вторая колонка
                                        for idx, row in enumerate(data):
                                            doc = DocxTemplate(name_file_template_doc)
                                            context = row
                                            doc.render(context)
                                            # Сохраняем файл
                                            name_file = f'{name_type_file} {row[name_main_column]} {row[name_second_column]}'
                                            name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                            threshold_name = 200 - (len(finish_path) + 10)
                                            if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                                raise OSError
                                            name_file = name_file[:threshold_name]  # ограничиваем название файла
                                            # Сохраняем файл
                                            save_result_file(finish_path, name_file, doc, idx, mode_pdf,name_os)
                                        zip_folder(finish_path,
                                                   f'Результаты тестирования {clean_first_name_folder}_{clean_second_name_folder}_{clean_third_name_folder}_{clean_four_name_folder}.zip')  # архивируем файлы docx
            else:
                if len(lst_number_column_folder_structure) == 1:
                    # Если нужно создавать одноуровневую структуру
                    # получаем название колонки
                    main_layer_name_column = df.columns[lst_number_column_folder_structure[0]]
                    lst_unique_value = df[main_layer_name_column].unique() # получаем список уникальных значений
                    for name_folder in lst_unique_value:
                        temp_df = df[df[main_layer_name_column] == name_folder] # фильтруем по названию
                        # Конвертируем датафрейм в список словарей
                        clean_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_', name_folder)  # очищаем название от лишних символов
                        finish_path = f'{path_to_end_folder_doc}/{clean_name_folder}'
                        if not os.path.exists(finish_path):
                            os.makedirs(finish_path)

                        # переименовываем колонки указанные в качестве идентифицирующих для того чтобы они отображалисьвнутри файла
                        if len(lst_number_column_name_file) == 1:
                            # если указана только одна колонка
                            name_column = temp_df.columns[lst_number_column_name_file[0]]
                            temp_df.rename(columns={name_column: 'Код_1'}, inplace=True)
                        elif len(lst_number_column_name_file) == 2:
                            name_main_column = temp_df.columns[lst_number_column_name_file[0]] # первая колонка
                            name_second_column = temp_df.columns[lst_number_column_name_file[1]] # вторая колонка
                            temp_df.rename(columns={name_main_column: 'Код_1',name_second_column: 'Код_2'}, inplace=True)

                        data = temp_df.to_dict('records')
                        # В зависимости от состояния чекбоксов обрабатываем файлы
                        # Создаем в цикле документы
                        if len(lst_number_column_name_file) == 1:
                            # если указана только одна колонка
                            name_column = temp_df.columns[lst_number_column_name_file[0]]
                            temp_df.rename(columns={name_column: 'Код_1'}, inplace=True)

                            for idx, row in enumerate(data):
                                doc = DocxTemplate(name_file_template_doc)
                                context = row
                                doc.render(context)
                                name_file = f'{name_type_file} {row[name_column]}'
                                name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                threshold_name = 200 - (len(finish_path) + 10)
                                if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                    raise OSError
                                name_file = name_file[:threshold_name]  # ограничиваем название файла
                                # Сохраняем файл
                                short_version_save_result_file(finish_path, name_file, doc, idx)


                        elif len(lst_number_column_name_file) == 2:
                            name_main_column = temp_df.columns[lst_number_column_name_file[0]]  # первая колонка
                            name_second_column = temp_df.columns[lst_number_column_name_file[1]]  # вторая колонка
                            for idx, row in enumerate(data):
                                doc = DocxTemplate(name_file_template_doc)
                                context = row
                                doc.render(context)
                                # Сохраняем файл
                                name_file = f'{name_type_file} {row[name_main_column]} {row[name_second_column]}'
                                name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                threshold_name = 200 - (len(finish_path) + 10)
                                if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                    raise OSError
                                name_file = name_file[:threshold_name]  # ограничиваем название файла
                                # Сохраняем файл
                                short_version_save_result_file(finish_path, name_file, doc, idx)
                elif len(lst_number_column_folder_structure) == 2:
                    # Если нужно создавать двухуровневую структуру
                    # получаем название колонки для первого уровня папок
                    name_first_layer_column = df.columns[lst_number_column_folder_structure[0]]
                    name_second_layer_column = df.columns[lst_number_column_folder_structure[1]]

                    lst_unique_value_first_layer = df[name_first_layer_column].unique()  # получаем список уникальных значений
                    for first_name_folder in lst_unique_value_first_layer:
                        clean_first_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                         first_name_folder)  # очищаем название от лишних символов

                        # получаем отфильтрованный датафрейм по значениям колонки первого уровня
                        temp_df_first_layer = df[df[name_first_layer_column] == first_name_folder]  # фильтруем по названию
                        lst_unique_value_second_layer = temp_df_first_layer[name_second_layer_column].unique()  # получаем список уникальных значений
                        # фильтруем по значениям колонки второго уровня
                        for second_name_folder in lst_unique_value_second_layer:
                            temp_df_second_layer = temp_df_first_layer[temp_df_first_layer[name_second_layer_column] == second_name_folder]
                            clean_second_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_', second_name_folder)  # очищаем название от лишних символов

                            finish_path = f'{path_to_end_folder_doc}/{clean_first_name_folder}/{clean_second_name_folder}'
                            if not os.path.exists(finish_path):
                                os.makedirs(finish_path)

                            # переименовываем колонки указанные в качестве идентифицирующих для того чтобы они отображалисьвнутри файла
                            if len(lst_number_column_name_file) == 1:
                                # если указана только одна колонка
                                name_column = temp_df_second_layer.columns[lst_number_column_name_file[0]]
                                temp_df_second_layer.rename(columns={name_column: 'Код_1'}, inplace=True)
                            elif len(lst_number_column_name_file) == 2:
                                name_main_column = temp_df_second_layer.columns[lst_number_column_name_file[0]]  # первая колонка
                                name_second_column = temp_df_second_layer.columns[lst_number_column_name_file[1]]  # вторая колонка
                                temp_df_second_layer.rename(columns={name_main_column: 'Код_1', name_second_column: 'Код_2'}, inplace=True)

                            data = temp_df_second_layer.to_dict('records') # конвертируем в список словарей
                            # Создаем в цикле документы
                            if len(lst_number_column_name_file) == 1:
                                # если указана только одна колонка
                                name_column = temp_df_second_layer.columns[lst_number_column_name_file[0]]
                                for idx, row in enumerate(data):
                                    doc = DocxTemplate(name_file_template_doc)
                                    context = row
                                    doc.render(context)
                                    # Сохраняем файл
                                    name_file = f'{name_type_file} {row[name_column]}'
                                    name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                    threshold_name = 200 - (len(finish_path) + 10)
                                    if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                        raise OSError
                                    name_file = name_file[:threshold_name]  # ограничиваем название файла
                                    # Сохраняем файл
                                    short_version_save_result_file(finish_path, name_file, doc, idx)

                            elif len(lst_number_column_name_file) == 2:
                                name_main_column = temp_df_second_layer.columns[lst_number_column_name_file[0]]  # первая колонка
                                name_second_column = temp_df_second_layer.columns[lst_number_column_name_file[1]]  # вторая колонка
                                for idx, row in enumerate(data):
                                    doc = DocxTemplate(name_file_template_doc)
                                    context = row
                                    doc.render(context)
                                    # Сохраняем файл
                                    name_file = f'{name_type_file} {row[name_main_column]} {row[name_second_column]}'
                                    name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                    threshold_name = 200 - (len(finish_path) + 10)
                                    if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                        raise OSError
                                    name_file = name_file[:threshold_name]  # ограничиваем название файла
                                    # Сохраняем файл
                                    short_version_save_result_file(finish_path, name_file, doc, idx)
                elif len(lst_number_column_folder_structure) == 3:
                    # Если нужно создавать трехуровневую структуру Например Школа-Класс-буква класса
                    # получаем названия колонок для трех уровней
                    name_first_layer_column = df.columns[lst_number_column_folder_structure[0]]
                    name_second_layer_column = df.columns[lst_number_column_folder_structure[1]]
                    name_third_layer_column = df.columns[lst_number_column_folder_structure[2]]

                    lst_unique_value_first_layer = df[name_first_layer_column].unique()  # получаем список уникальных значений
                    for first_name_folder in lst_unique_value_first_layer:
                        clean_first_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                         first_name_folder)  # очищаем название от лишних символов

                        # получаем отфильтрованный датафрейм по значениям колонки первого уровня
                        temp_df_first_layer = df[df[name_first_layer_column] == first_name_folder]  # фильтруем по названию
                        lst_unique_value_second_layer = temp_df_first_layer[
                            name_second_layer_column].unique()  # получаем список уникальных значений второго уровня
                        # фильтруем по значениям колонки второго уровня
                        for second_name_folder in lst_unique_value_second_layer:
                            temp_df_second_layer = temp_df_first_layer[
                                temp_df_first_layer[name_second_layer_column] == second_name_folder]
                            clean_second_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                              second_name_folder)  # очищаем название от лишних символов
                            lst_unique_value_third_layer = temp_df_second_layer[
                                name_third_layer_column].unique()  # получаем список уникальных значений третьего уровня
                            for third_name_folder in lst_unique_value_third_layer:
                                clean_third_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                                  third_name_folder)  # очищаем название от лишних символов
                                temp_df_third_layer = temp_df_second_layer[
                                    temp_df_second_layer[name_third_layer_column] == third_name_folder]

                                finish_path = f'{path_to_end_folder_doc}/{clean_first_name_folder}/{clean_second_name_folder}/{clean_third_name_folder}'
                                if not os.path.exists(finish_path):
                                    os.makedirs(finish_path)
                                if len(lst_number_column_name_file) == 1:
                                    # если указана только одна колонка
                                    name_column = temp_df_third_layer.columns[lst_number_column_name_file[0]]
                                    temp_df_third_layer.rename(columns={name_column: 'Код_1'}, inplace=True)
                                elif len(lst_number_column_name_file) == 2:
                                    name_main_column = temp_df_third_layer.columns[lst_number_column_name_file[0]]  # первая колонка
                                    name_second_column = temp_df_third_layer.columns[lst_number_column_name_file[1]]  # вторая колонка
                                    temp_df_third_layer.rename(columns={name_main_column: 'Код_1', name_second_column: 'Код_2'}, inplace=True)

                                data = temp_df_third_layer.to_dict('records')  # конвертируем в список словарей

                                # Создаем в цикле документы
                                if len(lst_number_column_name_file) == 1:
                                    # если указана только одна колонка
                                    name_column = temp_df_third_layer.columns[lst_number_column_name_file[0]]
                                    for idx, row in enumerate(data):
                                        doc = DocxTemplate(name_file_template_doc)
                                        context = row
                                        doc.render(context)
                                        # Сохраняем файл
                                        name_file = f'{name_type_file} {row[name_column]}'
                                        name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                        threshold_name = 200 - (len(finish_path) + 10)
                                        if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                            raise OSError
                                        name_file = name_file[:threshold_name]  # ограничиваем название файла
                                        # Сохраняем файл
                                        short_version_save_result_file(finish_path, name_file, doc, idx)

                                elif len(lst_number_column_name_file) == 2:
                                    name_main_column = temp_df_third_layer.columns[
                                        lst_number_column_name_file[0]]  # первая колонка
                                    name_second_column = temp_df_third_layer.columns[
                                        lst_number_column_name_file[1]]  # вторая колонка
                                    for idx, row in enumerate(data):
                                        doc = DocxTemplate(name_file_template_doc)
                                        context = row
                                        doc.render(context)
                                        # Сохраняем файл
                                        name_file = f'{name_type_file} {row[name_main_column]} {row[name_second_column]}'
                                        name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                        threshold_name = 200 - (len(finish_path) + 10)
                                        if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                            raise OSError
                                        name_file = name_file[:threshold_name]  # ограничиваем название файла
                                        # Сохраняем файл
                                        short_version_save_result_file(finish_path, name_file, doc, idx)
                elif len(lst_number_column_folder_structure) == 4:
                    # Если нужно создавать четырех уровневую структуру Например Результат- Число результата--Класс-буква класса
                    # получаем названия колонок для 4 уровней
                    name_first_layer_column = df.columns[lst_number_column_folder_structure[0]]
                    name_second_layer_column = df.columns[lst_number_column_folder_structure[1]]
                    name_third_layer_column = df.columns[lst_number_column_folder_structure[2]]
                    name_four_layer_column = df.columns[lst_number_column_folder_structure[3]]

                    lst_unique_value_first_layer = df[name_first_layer_column].unique()  # получаем список уникальных значений
                    for first_name_folder in lst_unique_value_first_layer:
                        clean_first_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                         first_name_folder)  # очищаем название от лишних символов

                        # получаем отфильтрованный датафрейм по значениям колонки первого уровня
                        temp_df_first_layer = df[df[name_first_layer_column] == first_name_folder]  # фильтруем по названию
                        lst_unique_value_second_layer = temp_df_first_layer[
                            name_second_layer_column].unique()  # получаем список уникальных значений второго уровня
                        # фильтруем по значениям колонки второго уровня
                        for second_name_folder in lst_unique_value_second_layer:
                            temp_df_second_layer = temp_df_first_layer[
                                temp_df_first_layer[name_second_layer_column] == second_name_folder]
                            clean_second_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                              second_name_folder)  # очищаем название от лишних символов
                            lst_unique_value_third_layer = temp_df_second_layer[
                                name_third_layer_column].unique()  # получаем список уникальных значений третьего уровня

                            for third_name_folder in lst_unique_value_third_layer:
                                clean_third_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                                  third_name_folder)  # очищаем название от лишних символов
                                temp_df_third_layer = temp_df_second_layer[
                                    temp_df_second_layer[name_third_layer_column] == third_name_folder]
                                lst_unique_value_four_layer = temp_df_third_layer[
                                    name_four_layer_column].unique()  # получаем список уникальных значений четвертого уровня

                                for four_name_folder in lst_unique_value_four_layer:
                                    clean_four_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                                     four_name_folder)  # очищаем название от лишних символов
                                    temp_df_four_layer = temp_df_third_layer[
                                        temp_df_third_layer[name_four_layer_column] == four_name_folder]


                                    finish_path = f'{path_to_end_folder_doc}/{clean_first_name_folder}/{clean_second_name_folder}/{clean_third_name_folder}/{clean_four_name_folder}'
                                    if not os.path.exists(finish_path):
                                        os.makedirs(finish_path)
                                    if len(lst_number_column_name_file) == 1:
                                        # если указана только одна колонка
                                        name_column = temp_df_third_layer.columns[lst_number_column_name_file[0]]
                                        temp_df_four_layer.rename(columns={name_column: 'Код_1'}, inplace=True)
                                    elif len(lst_number_column_name_file) == 2:
                                        name_main_column = temp_df_four_layer.columns[lst_number_column_name_file[0]]  # первая колонка
                                        name_second_column = temp_df_four_layer.columns[lst_number_column_name_file[1]]  # вторая колонка
                                        temp_df_four_layer.rename(columns={name_main_column: 'Код_1', name_second_column: 'Код_2'}, inplace=True)

                                    data = temp_df_four_layer.to_dict('records')  # конвертируем в список словарей

                                    # Создаем в цикле документы
                                    if len(lst_number_column_name_file) == 1:
                                        # если указана только одна колонка
                                        name_column = temp_df_four_layer.columns[lst_number_column_name_file[0]]
                                        for idx, row in enumerate(data):
                                            doc = DocxTemplate(name_file_template_doc)
                                            context = row
                                            doc.render(context)
                                            # Сохраняем файл
                                            name_file = f'{name_type_file} {row[name_column]}'
                                            name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                            threshold_name = 200 - (len(finish_path) + 10)
                                            if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                                raise OSError
                                            name_file = name_file[:threshold_name]  # ограничиваем название файла
                                            # Сохраняем файл
                                            short_version_save_result_file(finish_path, name_file, doc, idx)
                                    elif len(lst_number_column_name_file) == 2:
                                        name_main_column = temp_df_four_layer.columns[
                                            lst_number_column_name_file[0]]  # первая колонка
                                        name_second_column = temp_df_four_layer.columns[
                                            lst_number_column_name_file[1]]  # вторая колонка
                                        for idx, row in enumerate(data):
                                            doc = DocxTemplate(name_file_template_doc)
                                            context = row
                                            doc.render(context)
                                            # Сохраняем файл
                                            name_file = f'{name_type_file} {row[name_main_column]} {row[name_second_column]}'
                                            name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                            threshold_name = 200 - (len(finish_path) + 10)
                                            if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                                raise OSError
                                            name_file = name_file[:threshold_name]  # ограничиваем название файла
                                            # Сохраняем файл
                                            short_version_save_result_file(finish_path, name_file, doc, idx)

        # Удаляем файл с разрывом страницы
        try:
            os.remove(template_page_break_path_finish)
        except OSError as e:
            print("Ошибка при попытке удаления файла: {}".format(e))


    except NameError as e:
        messagebox.showerror('Лахеcис Обработка результатов профориентационных тестов',
                             f'Выберите шаблон,файл с данными и папку куда будут генерироваться файлы.')
    except KeyError as e:
        messagebox.showerror('Лахеcис Обработка результатов профориентационных тестов',
                             f'В таблице не найдена указанная колонка {e.args}.')
    except PermissionError:
        messagebox.showerror('Лахеcис Обработка результатов профориентационных тестов',
                             f'Закройте все файлы Word созданные Вестой.')
    except OSError:
        messagebox.showerror('Лахеcис Обработка результатов профориентационных тестов',
                             f'Слишком длинный путь к создаваемым файлам.\nПроверьте длину текста в колонках которые вы выбрали\n'
                             f'для создания структуры папки, названия файла, типа файла.\n'
                             f'Попробуйте перенести итоговую папку в корень диска.'
                             )

    except NotNumberColumn:
        messagebox.showerror('Лахеcис Обработка результатов профориентационных тестов',
                             f'Колонки с таким порядковым номером нет в таблице.\nПроверьте правильность введенных данных.'
                             )
    except NoMoreNumberColumn:
        messagebox.showerror('Лахеcис Обработка результатов профориентационных тестов',
                             f'Проверьте количество введенных чисел на шаге 3 или шаге 4.\n'
                             f'Для шага 3 (структура папок) не более 4 чисел разделенных запятыми.\n'
                             f'Например 3,5,12,14.\n'
                             f'Для шага 4 (названия файлов) не более 2 чисел разделенных запятыми.\n'
                             f'Например 6,7.'
                             )
    except PdfLinux:
        messagebox.showerror('Лахеcис Обработка результатов профориентационных тестов',
                             f'Создание PDF файлов возможно только в Windows! Снимите галочки с создания PDF документов И поставьте галочку для включения полного режима'
                             )
    except PDFAndFull:
        messagebox.showerror('Лахеcис Обработка результатов профориентационных тестов',
                             f'Уберите галочку из одной опции. Запуск программы возможен только если стоит только одна галочка'
                             )
    else:
        messagebox.showinfo('Лахеcис Обработка результатов профориентационных тестов',
                            'Создание документов завершено!')


if __name__ == '__main__':
    main_name_file_data_doc = 'c:/Users/1/PycharmProjects/Lachesis/data/Таблица с обезличенными результатами.xlsx'
    main_name_file_template_doc = 'c:/Users/1/PycharmProjects/Lachesis/data/Шаблон Отчет о результатах комплексного профориентационного тестирования.docx'
    main_path_to_end_folder_doc = 'c:/Users/1/PycharmProjects/Lachesis/data/Результат'
    main_folder_structure = '9,3,4'
    main_name_file = '5,6'
    main_name_type_file = 'Результат тестирования'
    main_mode_pdf = 'No'
    main_mode_full = 'Yes'

    generate_result_docs(main_name_file_data_doc,main_name_file_template_doc,main_path_to_end_folder_doc,
                         main_folder_structure,main_name_file,main_name_type_file,main_mode_pdf,main_mode_full)

    print('Lindy Booth')
